import concurrent.futures
import functools
import hashlib
import json
import logging
import os
import re
import tempfile
import threading
import time
import unicodedata
from datetime import date
from pathlib import Path

log = logging.getLogger(__name__)

from dotenv import load_dotenv
from flask import Flask, request, jsonify, render_template, Response, session

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")


def _normalize_search(s: str) -> str:
    """Normaliza string para busca: remove acentos, pontuação e caixa."""
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    s = s.lower()
    s = re.sub(r"[^a-z0-9 ]", " ", s)
    return re.sub(r"\s+", " ", s).strip()


from datetime import timedelta
from werkzeug.middleware.proxy_fix import ProxyFix

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-insecure-key-troque-no-env")
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = os.environ.get("RENDER", "") != ""
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=90)

CIFRAS_ROOT = os.environ.get("CIFRAS_ROOT", str(Path.home() / "OneDrive" / "Cifras"))
CIFRAS_FOLDER_ID = os.environ.get("CIFRAS_FOLDER_ID", "")
SUPPORTED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt", ".md"}

# E-mails com acesso de owner (separados por vírgula)
_OWNER_EMAILS = {
    e.strip().lower()
    for e in os.environ.get("OWNER_EMAIL", "").split(",")
    if e.strip()
}

# Registra blueprint de autenticação
from auth import bp as auth_bp, login_required, get_service, current_user, is_oauth_configured
app.register_blueprint(auth_bp)


# ---------------------------------------------------------------------------
# Controle de acesso por role
# ---------------------------------------------------------------------------

def is_owner() -> bool:
    """True se o usuário logado é owner (ou se OWNER_EMAIL não está definido)."""
    if not _OWNER_EMAILS:
        return True
    return current_user().get("email", "").lower().strip() in _OWNER_EMAILS


def owner_required(f):
    """Decorator: bloqueia a rota para não-owners com 403."""
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if not is_owner():
            return jsonify({"error": "Permissão negada"}), 403
        return f(*args, **kwargs)
    return decorated


def _get_user_data_folder_id(svc):
    """Owner usa CIFRAS_FOLDER_ID (compatibilidade). Outros: pasta própria no Drive."""
    if is_owner() or not _use_drive():
        return CIFRAS_FOLDER_ID
    import drive as drv
    return drv.get_user_data_folder(svc)


# ---------------------------------------------------------------------------
# Service Worker
# ---------------------------------------------------------------------------

@app.route("/sw.js")
def service_worker():
    from flask import send_from_directory
    resp = send_from_directory("static", "sw.js")
    resp.headers["Service-Worker-Allowed"] = "/"
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    return resp


# ---------------------------------------------------------------------------
# Modo de operação
# ---------------------------------------------------------------------------

def _use_drive():
    """Drive ativo se OAuth configurado E CIFRAS_FOLDER_ID definido."""
    return is_oauth_configured() and bool(CIFRAS_FOLDER_ID)


_AUTH_ERROR_SIGNALS = (
    "invalid_grant", "token has been expired", "token has been revoked",
    "unauthorized", "401", "revoked", "invalid_token",
)

def _is_auth_error(e):
    full = (str(e) + " " + str(getattr(e, "__cause__", "") or "")).lower()
    return any(s in full for s in _AUTH_ERROR_SIGNALS)

def _auth_error_response():
    session.clear()
    return jsonify({"error": "Sessão expirada", "login_url": "/"}), 401


# ---------------------------------------------------------------------------
# Extração de texto — por bytes
# ---------------------------------------------------------------------------

def extract_text_from_bytes(content_bytes, extension):
    ext = extension.lower()
    if ext in (".docx", ".doc"):
        return _docx_from_bytes(content_bytes)
    elif ext == ".pdf":
        return _pdf_from_bytes(content_bytes)
    elif ext in (".txt", ".md"):
        text = content_bytes.decode("utf-8", errors="replace")
        return _strip_frontmatter(text)
    return "[Formato não suportado]"


def _docx_from_bytes(content_bytes):
    try:
        import io as _io
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(_io.BytesIO(content_bytes))
        lines = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Parágrafo normal
                text = "".join(n.text or "" for n in element.iter() if n.tag.endswith("}t"))
                lines.append(text)

            elif tag == "tbl":
                # Tabela — lê cada linha juntando células com espaço
                for row in element.findall(".//" + qn("w:tr")):
                    cells = []
                    for cell in row.findall(".//" + qn("w:tc")):
                        cell_text = "".join(
                            n.text or "" for n in cell.iter() if n.tag.endswith("}t")
                        )
                        if cell_text.strip():
                            cells.append(cell_text.strip())
                    if cells:
                        lines.append("   ".join(cells))

        return "\n".join(lines)
    except Exception:
        return _docx_via_com(content_bytes)


def _docx_via_com(content_bytes):
    tmp_path = None
    try:
        import win32com.client
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(content_bytes)
            tmp_path = tmp.name
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(str(Path(tmp_path).resolve()), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return text
    except ImportError:
        return "[Erro: .doc binário não suportado. Instale pywin32.]"
    except Exception as e:
        return f"[Erro ao ler via Word COM: {e}]"
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


def _pdf_from_bytes(content_bytes):
    try:
        import io as _io
        import fitz
        doc = fitz.open(stream=_io.BytesIO(content_bytes), filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        return f"[Erro ao ler .pdf: {e}]"


def _strip_frontmatter(text):
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            return parts[2].strip()
    return text


def _parse_frontmatter(text):
    """Retorna (body, meta_dict) extraindo YAML frontmatter simples."""
    meta = {"artist": "", "key": "", "title": "", "tags": [], "capo": "", "youtube": ""}
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            for line in parts[1].splitlines():
                if ":" in line:
                    k, _, v = line.partition(":")
                    key = k.strip().lower()
                    val = v.strip().strip('"').strip("'")
                    if key == "tags":
                        # Parse simple YAML list: [a, b, c] or a, b, c
                        val = val.strip("[]")
                        meta["tags"] = [t.strip() for t in val.split(",") if t.strip()] if val else []
                    else:
                        meta[key] = val
            return parts[2].strip(), meta
    return text, meta


# ---------------------------------------------------------------------------
# Extração de texto — por caminho local
# ---------------------------------------------------------------------------

def extract_text(path):
    try:
        content_bytes = Path(path).read_bytes()
        return extract_text_from_bytes(content_bytes, Path(path).suffix)
    except Exception as e:
        return f"[Erro ao ler arquivo: {e}]"


# ---------------------------------------------------------------------------
# Biblioteca local
# ---------------------------------------------------------------------------

def _md_meta(path):
    """Lê apenas o frontmatter de um .md sem carregar o corpo inteiro."""
    try:
        with open(path, encoding="utf-8", errors="replace") as fh:
            head = fh.read(1024)
        _, meta = _parse_frontmatter(head)
        return meta
    except Exception:
        return {"artist": "", "key": "", "title": ""}


def _song_entry(f, section, category, read_meta=False):
    entry = {"name": f.stem, "path": str(f), "section": section, "category": category,
             "artist": "", "key": ""}
    if read_meta and f.suffix.lower() == ".md":
        meta = _md_meta(f)
        entry["artist"] = meta.get("artist", "")
        entry["key"] = meta.get("key", "")
        if meta.get("title"):
            entry["name"] = meta["title"]
    return entry


def scan_library_local():
    root = Path(CIFRAS_ROOT)
    library = {}
    if not root.exists():
        return library
    for section_dir in sorted(root.iterdir()):
        if not section_dir.is_dir():
            continue
        sname = section_dir.name
        library[sname] = {}
        for item in sorted(section_dir.iterdir()):
            if item.is_dir():
                songs = _collect_local(item, sname, item.name)
                if songs:
                    library[sname][item.name] = songs
            elif item.suffix.lower() in SUPPORTED_EXTENSIONS:
                library[sname].setdefault("_raiz", [])
                library[sname]["_raiz"].append(_song_entry(item, sname, "_raiz", read_meta=False))
    return library


def _collect_local(folder, section, category):
    return [
        _song_entry(f, section, category, read_meta=False)
        for f in sorted(folder.iterdir())
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


def flatten_songs(library):
    return [song for cats in library.values() for songs in cats.values() for song in songs]


# ---------------------------------------------------------------------------
# Repertórios (persistidos em arquivo JSON)
# ---------------------------------------------------------------------------

REPERTORIOS_LOCAL = Path(__file__).parent / "_repertorios.json"
_rep_lock = threading.Lock()
_reps_cache: dict = {}   # { email: {"data": ..., "file_id": ...} }

def _load_reps_local():
    try:
        return json.loads(REPERTORIOS_LOCAL.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _save_reps_local(data):
    REPERTORIOS_LOCAL.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def _load_reps():
    if _use_drive():
        import drive as drv
        svc = get_service()
        folder_id = _get_user_data_folder_id(svc)
        email = current_user().get("email", "_local")
        # Always read fresh from Drive — cache only the file_id for fast writes
        data, file_id = drv.load_repertorios(svc, folder_id)
        cached = _reps_cache.get(email, {})
        _reps_cache[email] = {"data": None, "file_id": file_id or cached.get("file_id")}
        return data
    return _load_reps_local()

def _save_reps(data):
    if _use_drive():
        import drive as drv
        svc = get_service()
        folder_id = _get_user_data_folder_id(svc)
        email = current_user().get("email", "_local")
        cached = _reps_cache.get(email)
        if cached and cached.get("file_id"):
            file_id = cached["file_id"]
        else:
            _, file_id = drv.load_repertorios(svc, folder_id)
        drv.save_repertorios(svc, file_id, data)
        _reps_cache[email] = {"data": data, "file_id": file_id}
    else:
        _save_reps_local(data)


# ─── Grupos ──────────────────────────────────────────────────────────────────

_groups_cache: dict = {}  # { email: {"file_id": str} }

def _load_groups():
    if _use_drive():
        import drive as drv
        svc = get_service()
        folder_id = _get_user_data_folder_id(svc)
        email = current_user().get("email", "_local")
        data, file_id = drv.load_groups(svc, folder_id)
        _groups_cache[email] = {"file_id": file_id}
        return data
    return {}

def _save_groups(data):
    if _use_drive():
        import drive as drv
        svc = get_service()
        folder_id = _get_user_data_folder_id(svc)
        email = current_user().get("email", "_local")
        cached = _groups_cache.get(email)
        if cached and cached.get("file_id"):
            file_id = cached["file_id"]
        else:
            _, file_id = drv.load_groups(svc, folder_id)
        drv.save_groups(svc, file_id, data)
        _groups_cache[email] = {"file_id": file_id}


@app.route("/api/groups", methods=["GET"])
@login_required
def api_groups_list():
    try:
        return jsonify(_load_groups())
    except Exception as e:
        if _is_auth_error(e):
            return _auth_error_response()
        return jsonify({}), 200


@app.route("/api/groups", methods=["POST"])
@login_required
def api_groups_create():
    body = request.get_json(force=True)
    name = (body.get("name") or "").strip()
    members = [e.strip().lower() for e in body.get("members", []) if e.strip()]
    if not name:
        return jsonify({"error": "Nome obrigatório"}), 400
    groups = _load_groups()
    if len(groups) >= 20:
        return jsonify({"error": "Limite de 20 grupos atingido"}), 400
    gid = "grp_" + os.urandom(6).hex()
    from datetime import datetime
    group = {"id": gid, "name": name, "members": members,
             "created_at": datetime.utcnow().isoformat()}
    groups[gid] = group
    _save_groups(groups)
    return jsonify(group), 201


@app.route("/api/groups/<gid>", methods=["PUT"])
@login_required
def api_groups_update(gid):
    from datetime import datetime
    body = request.get_json(force=True)
    groups = _load_groups()
    if gid not in groups:
        return jsonify({"error": "Grupo não encontrado"}), 404
    old_members = set(groups[gid].get("members", []))
    if "name" in body:
        name = body["name"].strip()
        if name:
            groups[gid]["name"] = name
    if "members" in body:
        groups[gid]["members"] = [e.strip().lower() for e in body["members"] if e.strip()]
    new_members = set(groups[gid]["members"])
    added = new_members - old_members
    _save_groups(groups)

    # Propaga compartilhamentos existentes do grupo para membros novos
    if added:
        me = current_user()
        my_email = me.get("email", "").lower()
        reps = _load_reps()
        with _shares_lock:
            shares = _load_shares_raw()
            group_rep_ids = {s["rep_id"] for s in shares.values()
                             if s.get("group_id") == gid
                             and s.get("from_email", "").lower() == my_email}
            for email in added:
                if email == my_email:
                    continue
                for rep_id in group_rep_ids:
                    _do_share_rep(rep_id, email, me, reps, shares, group_id=gid)
            if group_rep_ids:
                _save_shares_raw(shares)

    return jsonify(groups[gid])


@app.route("/api/groups/<gid>", methods=["DELETE"])
@login_required
def api_groups_delete(gid):
    groups = _load_groups()
    if gid not in groups:
        return jsonify({"error": "Grupo não encontrado"}), 404
    del groups[gid]
    _save_groups(groups)
    return jsonify({"ok": True})


@app.route("/api/repertorios", methods=["GET"])
@login_required
def api_reps_list():
    try:
        return jsonify(_load_reps())
    except Exception as e:
        if _is_auth_error(e):
            return _auth_error_response()
        return jsonify({"error": str(e)}), 500


@app.route("/api/repertorios", methods=["POST"])
@login_required
def api_reps_create():
    from datetime import datetime
    data = request.get_json(force=True)
    name = data.get("name", "").strip()
    songs = data.get("songs", [])
    if not name:
        return jsonify({"error": "Nome obrigatório"}), 400
    rep_id = "rpt_" + os.urandom(6).hex()
    now = datetime.utcnow().isoformat()
    rep = {"id": rep_id, "name": name, "songs": songs, "created_at": now, "updated_at": now}
    with _rep_lock:
        reps = _load_reps()
        if len(reps) >= 5:
            return jsonify({"error": "Limite de 5 repertórios atingido"}), 400
        reps[rep_id] = rep
        _save_reps(reps)
    return jsonify(rep), 201


@app.route("/api/repertorios/<rep_id>", methods=["PUT"])
@login_required
def api_reps_update(rep_id):
    from datetime import datetime
    data = request.get_json(force=True)
    with _rep_lock:
        reps = _load_reps()
        if rep_id not in reps:
            return jsonify({"error": "Não encontrado"}), 404
        if "name" in data:
            reps[rep_id]["name"] = data["name"].strip() or reps[rep_id]["name"]
        if "songs" in data:
            reps[rep_id]["songs"] = data["songs"]
        reps[rep_id]["updated_at"] = datetime.utcnow().isoformat()
        _save_reps(reps)

    # Live sharing: propaga alterações de songs para todos os shares ativos deste rep
    if "songs" in data or "name" in data:
        try:
            from_email = current_user().get("email", "").lower()
            new_songs = reps[rep_id].get("songs", [])
            new_name  = reps[rep_id].get("name", "")
            with _shares_lock:
                shares = _load_shares_raw()
                changed = False
                for share in shares.values():
                    if share.get("rep_id") == rep_id and \
                       share.get("from_email", "").lower() == from_email:
                        if "songs" in data:
                            share["songs"] = new_songs
                        if "name" in data:
                            share["rep_name"] = new_name
                        changed = True
                if changed:
                    _save_shares_raw(shares)
        except Exception:
            pass  # não bloqueia a resposta principal

    return jsonify(reps[rep_id])


@app.route("/api/repertorios/<rep_id>", methods=["DELETE"])
@login_required
def api_reps_delete(rep_id):
    with _rep_lock:
        reps = _load_reps()
        if rep_id not in reps:
            return jsonify({"error": "Não encontrado"}), 404
        del reps[rep_id]
        _save_reps(reps)
    # Remove todas as shares deste rep
    my_email = current_user().get("email", "").lower()
    with _shares_lock:
        shares = _load_shares_raw()
        before = len(shares)
        shares = {k: v for k, v in shares.items()
                  if not (v.get("rep_id") == rep_id and
                          v.get("from_email", "").lower() == my_email)}
        if len(shares) < before:
            _save_shares_raw(shares)
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Compartilhamento de repertórios
# ---------------------------------------------------------------------------

SHARES_LOCAL = Path(__file__).parent / "_shares.json"
_shares_lock = threading.Lock()
_shares_drive_file_id: str | None = None  # cache do file_id no Drive


def _load_shares_raw():
    """Lê o registro central de shares.
    Sempre lê do arquivo local (fonte de verdade compartilhada entre workers).
    Cold start: carrega do Drive e grava o arquivo local."""
    global _shares_drive_file_id
    try:
        data = json.loads(SHARES_LOCAL.read_text(encoding="utf-8"))
        log.debug("[shares] lido arquivo local (%d shares)", len(data))
        return data
    except Exception as e:
        log.info("[shares] arquivo local indisponível (%s) — tentando Drive", e)
    # Cold start: arquivo local não existe ainda, carrega do Drive
    if _use_drive() and CIFRAS_FOLDER_ID:
        try:
            import drive as drv
            svc = get_service()
            data, fid = drv.load_shares(svc, CIFRAS_FOLDER_ID)
            if fid:
                _shares_drive_file_id = fid
            log.info("[shares] carregado do Drive (%d shares)", len(data))
            try:
                SHARES_LOCAL.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
                log.info("[shares] arquivo local criado em %s", SHARES_LOCAL)
            except Exception as e:
                log.error("[shares] falha ao gravar arquivo local: %s", e)
            return data
        except Exception as e:
            log.error("[shares] falha ao carregar do Drive: %s", e)
    return {}


def _save_shares_raw(data):
    """Persiste o registro de shares."""
    global _shares_drive_file_id
    try:
        SHARES_LOCAL.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        log.info("[shares] arquivo local salvo (%d shares)", len(data))
    except Exception as e:
        log.error("[shares] falha ao gravar arquivo local: %s", e)
    # Só o owner tem permissão de escrita no Drive — viewer usa apenas arquivo local
    if _use_drive() and CIFRAS_FOLDER_ID and is_owner():
        try:
            import drive as drv
            svc = get_service()
            if not _shares_drive_file_id:
                _, _shares_drive_file_id = drv.load_shares(svc, CIFRAS_FOLDER_ID)
            drv.save_shares(svc, _shares_drive_file_id, data)
            log.info("[shares] Drive atualizado")
        except Exception as e:
            log.error("[shares] falha ao salvar no Drive: %s", e)


def _do_share_rep(rep_id, to_email, me, reps, shares, group_id=None):
    """Cria um share individual. Retorna (share_dict | None, error_str | None)."""
    from datetime import datetime
    rep = reps.get(rep_id)
    if not rep:
        return None, "Repertório não encontrado"
    my_email = me.get("email", "").lower()
    if to_email == my_email:
        return None, "Você não pode compartilhar com você mesmo"
    for s in shares.values():
        if (s.get("rep_id") == rep_id
                and s.get("from_email", "").lower() == my_email
                and s.get("to_email", "").lower() == to_email):
            return None, f"Já compartilhado com {to_email}"
    share_id = "shr_" + os.urandom(6).hex()
    share = {
        "id": share_id, "rep_id": rep_id, "rep_name": rep["name"],
        "songs": rep.get("songs", []),
        "from_email": my_email, "from_name": me.get("name", my_email),
        "from_picture": me.get("picture", ""), "to_email": to_email,
        "shared_at": datetime.utcnow().isoformat(),
        "seen_by": [], "dismissed_by": [],
    }
    if group_id:
        share["group_id"] = group_id
    shares[share_id] = share
    return share, None


@app.route("/api/share-rep", methods=["POST"])
@login_required
def api_share_rep():
    body = request.get_json(force=True)
    rep_id = (body.get("rep_id") or "").strip()
    to_email = (body.get("to_email") or "").strip().lower()
    group_id = (body.get("group_id") or "").strip()
    if not rep_id or (not to_email and not group_id):
        return jsonify({"error": "rep_id e to_email (ou group_id) são obrigatórios"}), 400
    reps = _load_reps()
    me = current_user()
    my_email = me.get("email", "").lower()

    # Compartilhamento para grupo
    if group_id:
        groups = _load_groups()
        group = groups.get(group_id)
        if not group:
            return jsonify({"error": "Grupo não encontrado"}), 404
        emails = [e for e in group.get("members", []) if e and e != my_email]
        if not emails:
            return jsonify({"error": "Grupo sem membros válidos"}), 400
        from datetime import datetime
        created = []
        with _shares_lock:
            shares = _load_shares_raw()
            rep = reps.get(rep_id)
            for email in emails:
                # Se já existe share para este membro, atualiza (renova songs e timestamp)
                existing = next((s for s in shares.values()
                                 if s.get("rep_id") == rep_id
                                 and s.get("from_email", "").lower() == my_email
                                 and s.get("to_email", "").lower() == email), None)
                if existing:
                    existing["songs"] = rep.get("songs", []) if rep else existing["songs"]
                    existing["shared_at"] = datetime.utcnow().isoformat()
                    created.append(existing)
                else:
                    share, _ = _do_share_rep(rep_id, email, me, reps, shares, group_id=group_id)
                    if share:
                        created.append(share)
            _save_shares_raw(shares)
        return jsonify({"shared": len(created), "skipped": len(emails) - len(created)}), 201

    # Compartilhamento individual
    if not to_email:
        return jsonify({"error": "rep_id e to_email são obrigatórios"}), 400
    reps = _load_reps()
    rep = reps.get(rep_id)
    if not rep:
        return jsonify({"error": "Repertório não encontrado"}), 404
    if to_email == my_email:
        return jsonify({"error": "Você não pode compartilhar com você mesmo"}), 400
    with _shares_lock:
        shares = _load_shares_raw()
        for s in shares.values():
            if (s.get("rep_id") == rep_id
                    and s.get("from_email", "").lower() == my_email
                    and s.get("to_email", "").lower() == to_email):
                return jsonify({"error": "Já compartilhado com este usuário"}), 400
        share_id = "shr_" + os.urandom(6).hex()
        from datetime import datetime
        share = {
            "id": share_id,
            "rep_id": rep_id,
            "rep_name": rep["name"],
            "songs": rep.get("songs", []),
            "from_email": my_email,
            "from_name": me.get("name", my_email),
            "from_picture": me.get("picture", ""),
            "to_email": to_email,
            "shared_at": datetime.utcnow().isoformat(),
            "seen_by": [],
            "dismissed_by": [],
        }
        shares[share_id] = share
        _save_shares_raw(shares)
    return jsonify(share), 201


@app.route("/api/shares-by-me", methods=["GET"])
@login_required
def api_shares_by_me():
    my_email = current_user().get("email", "").lower()
    shares = _load_shares_raw()
    result = [s for s in shares.values() if s.get("from_email", "").lower() == my_email]
    return jsonify(result)


@app.route("/api/shared-with-me", methods=["GET"])
@login_required
def api_shared_with_me():
    my_email = current_user().get("email", "").lower()
    shares = _load_shares_raw()
    result = []
    for s in shares.values():
        if s.get("to_email", "").lower() != my_email:
            continue
        if my_email in [e.lower() for e in s.get("dismissed_by", [])]:
            continue
        result.append(s)
    result.sort(key=lambda x: x.get("shared_at", ""), reverse=True)
    return jsonify(result)


@app.route("/api/share/<share_id>", methods=["DELETE"])
@login_required
def api_unshare(share_id):
    my_email = current_user().get("email", "").lower()
    with _shares_lock:
        shares = _load_shares_raw()
        s = shares.get(share_id)
        if not s:
            return jsonify({"error": "Não encontrado"}), 404
        if s.get("from_email", "").lower() != my_email:
            return jsonify({"error": "Permissão negada"}), 403
        del shares[share_id]
        _save_shares_raw(shares)
    return jsonify({"ok": True})


@app.route("/api/share/<share_id>/dismiss", methods=["POST"])
@login_required
def api_dismiss_share(share_id):
    my_email = current_user().get("email", "").lower()
    with _shares_lock:
        shares = _load_shares_raw()
        s = shares.get(share_id)
        if not s:
            return jsonify({"error": "Não encontrado"}), 404
        if s.get("to_email", "").lower() != my_email:
            return jsonify({"error": "Permissão negada"}), 403
        dismissed = [e.lower() for e in s.get("dismissed_by", [])]
        if my_email not in dismissed:
            s.setdefault("dismissed_by", []).append(my_email)
        _save_shares_raw(shares)
    return jsonify({"ok": True})


@app.route("/api/share/<share_id>/seen", methods=["POST"])
@login_required
def api_mark_share_seen(share_id):
    my_email = current_user().get("email", "").lower()
    with _shares_lock:
        shares = _load_shares_raw()
        s = shares.get(share_id)
        if not s:
            return jsonify({"error": "Não encontrado"}), 404
        if s.get("to_email", "").lower() != my_email:
            return jsonify({"error": "Permissão negada"}), 403
        seen = [e.lower() for e in s.get("seen_by", [])]
        if my_email not in seen:
            s.setdefault("seen_by", []).append(my_email)
        _save_shares_raw(shares)
    return jsonify({"ok": True})


@app.route("/api/notifications/count", methods=["GET"])
@login_required
def api_notif_count():
    my_email = current_user().get("email", "").lower()
    shares = _load_shares_raw()
    count = sum(
        1 for s in shares.values()
        if s.get("to_email", "").lower() == my_email
        and my_email not in [e.lower() for e in s.get("seen_by", [])]
        and my_email not in [e.lower() for e in s.get("dismissed_by", [])]
    )
    return jsonify({"count": count})


# ---------------------------------------------------------------------------
# Views tracking (persistido em views.json com lock para concorrência)
# ---------------------------------------------------------------------------

_views_lock = threading.Lock()
_views_cache: dict = {}   # { email: {"data": ..., "file_id": ...} }

def _load_views():
    """Carrega views do Drive para o usuário atual (com cache em memória)."""
    email = current_user().get("email", "_local")
    cached = _views_cache.get(email)
    if cached and cached.get("data") is not None:
        return cached["data"]
    if _use_drive():
        try:
            import drive as _drive
            svc = get_service()
            folder_id = _get_user_data_folder_id(svc)
            data, file_id = _drive.load_views(svc, folder_id)
            _views_cache[email] = {"data": data, "file_id": file_id}
            return data
        except Exception:
            return {}
    try:
        p = Path(__file__).parent / "views.json"
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _increment_view(key: str) -> int:
    """Incrementa o contador de uma música para o usuário atual e persiste."""
    email = current_user().get("email", "_local")
    with _views_lock:
        if _use_drive():
            try:
                import drive as _drive
                svc = get_service()
                folder_id = _get_user_data_folder_id(svc)
                cached = _views_cache.get(email)
                if not cached or cached.get("file_id") is None:
                    data, file_id = _drive.load_views(svc, folder_id)
                    _views_cache[email] = {"data": data, "file_id": file_id}
                views = _views_cache[email]["data"] or {}
                views[key] = views.get(key, 0) + 1
                _views_cache[email]["data"] = views
                _drive.save_views(svc, _views_cache[email]["file_id"], views)
                return views[key]
            except Exception as e:
                app.logger.error("Erro ao salvar view no Drive: %s", e)
                return 0
        else:
            p = Path(__file__).parent / "views.json"
            try:
                views = json.loads(p.read_text(encoding="utf-8"))
            except Exception:
                views = {}
            views[key] = views.get(key, 0) + 1
            p.write_text(json.dumps(views, ensure_ascii=False), encoding="utf-8")
            return views[key]

def _song_key(data):
    return data.get("fileId") or data.get("path", "")


# ---------------------------------------------------------------------------
# Preferências de tons (persistido em _preferences.json no Drive)
# ---------------------------------------------------------------------------

_prefs_lock = threading.Lock()
_prefs_cache: dict = {}   # { email: {"data": ..., "file_id": ...} }

VALID_SLOTS = {"my_key", "original_key", "alt_key", "my_capo"}


def _load_prefs():
    email = current_user().get("email", "_local")
    if _use_drive():
        try:
            import drive as _drive
            svc = get_service()
            folder_id = _get_user_data_folder_id(svc)
            # Always read fresh from Drive — cache only file_id for fast writes
            data, file_id = _drive.load_preferences(svc, folder_id)
            with _prefs_lock:
                cached = _prefs_cache.get(email, {})
                _prefs_cache[email] = {"data": None, "file_id": file_id or cached.get("file_id")}
            return data
        except Exception:
            pass
    return {}


def _save_prefs(data):
    import drive as _drive
    email = current_user().get("email", "_local")
    with _prefs_lock:
        if _use_drive():
            svc = get_service()
            cached = _prefs_cache.get(email)
            if not cached or cached.get("file_id") is None:
                folder_id = _get_user_data_folder_id(svc)
                _, file_id = _drive.load_preferences(svc, folder_id)
                if email not in _prefs_cache:
                    _prefs_cache[email] = {}
                _prefs_cache[email]["file_id"] = file_id
            _drive.save_preferences(svc, _prefs_cache[email]["file_id"], data)
        if email not in _prefs_cache:
            _prefs_cache[email] = {}
        _prefs_cache[email]["data"] = data


@app.route("/api/preferences")
@login_required
def api_get_preferences():
    """Retorna todas as preferências de tons ou as de uma música específica."""
    file_id = request.args.get("fileId", "")
    try:
        prefs = _load_prefs()
        if file_id:
            return jsonify(prefs.get(file_id, {}))
        return jsonify(prefs)
    except Exception as e:
        if _is_auth_error(e):
            return _auth_error_response()
        return jsonify({"error": str(e)}), 500


@app.route("/api/preferences", methods=["POST"])
@login_required
def api_save_preference():
    """Salva/atualiza um tom em um slot. Body: { fileId, slot, key }"""
    data = request.get_json() or {}
    file_id = data.get("fileId", "").strip()
    slot = data.get("slot", "").strip()
    key = data.get("key", "").strip()
    if not file_id or slot not in VALID_SLOTS:
        return jsonify({"error": "fileId e slot válido são obrigatórios"}), 400
    if slot != "my_capo" and not key:
        return jsonify({"error": "key é obrigatório para este slot"}), 400
    try:
        prefs = dict(_load_prefs())
        song_prefs = dict(prefs.get(file_id, {}))
        if slot == "my_capo" and (not key or key == "0"):
            song_prefs.pop("my_capo", None)
        else:
            song_prefs[slot] = key
        prefs[file_id] = song_prefs
        _save_prefs(prefs)
        return jsonify(song_prefs)
    except Exception as e:
        if _is_auth_error(e):
            return _auth_error_response()
        return jsonify({"error": str(e)}), 500


@app.route("/api/preferences", methods=["DELETE"])
@login_required
def api_delete_preference():
    """Remove um slot de tom. Body: { fileId, slot } ou { fileId } para remover tudo."""
    data = request.get_json() or {}
    file_id = data.get("fileId", "").strip()
    slot = data.get("slot", "").strip()
    if not file_id:
        return jsonify({"error": "fileId é obrigatório"}), 400
    try:
        prefs = dict(_load_prefs())
        if file_id not in prefs:
            return jsonify({})
        song_prefs = dict(prefs[file_id])
        if slot and slot in VALID_SLOTS:
            song_prefs.pop(slot, None)
        else:
            song_prefs = {}
        if song_prefs:
            prefs[file_id] = song_prefs
        else:
            prefs.pop(file_id, None)
        _save_prefs(prefs)
        return jsonify(song_prefs)
    except Exception as e:
        if _is_auth_error(e):
            return _auth_error_response()
        return jsonify({"error": str(e)}), 500


# ---------------------------------------------------------------------------
# Cache de biblioteca (evita re-scan a cada request)
# ---------------------------------------------------------------------------

_library_cache = {"data": None, "ts": 0}
_cache_lock = threading.Lock()
CACHE_TTL = 120  # segundos — re-escaneia se a última leitura foi há mais de 2 min

# Cache global de metadados por fileId — fonte de verdade compartilhada entre todos os usuários.
# Persistido em _songs_meta.json no CIFRAS_FOLDER_ID do Drive.
_songs_meta: dict = {}  # { fileId: {artist, key, capo, youtube} }
_songs_meta_file_id: str | None = None
_meta_lock = threading.Lock()

def _load_songs_meta_from_drive():
    """Carrega _songs_meta do Drive na inicialização do processo."""
    global _songs_meta, _songs_meta_file_id
    if not _use_drive() or not CIFRAS_FOLDER_ID:
        return
    try:
        import drive as drv
        svc = get_service()
        data, fid = drv.load_songs_meta(svc, CIFRAS_FOLDER_ID)
        with _meta_lock:
            _songs_meta = data
            _songs_meta_file_id = fid
    except Exception:
        pass

def _persist_songs_meta(svc=None):
    """Salva _songs_meta no Drive em background, usando o svc já autenticado.

    O svc DEVE ser passado pelo chamador (contexto de request). Chamar
    get_service() dentro da thread de background falha silenciosamente porque
    a Flask session não existe fora do contexto de request.
    """
    if svc is None or not _use_drive() or not CIFRAS_FOLDER_ID:
        return

    def _save(svc):
        global _songs_meta_file_id
        try:
            import drive as drv
            with _meta_lock:
                data = dict(_songs_meta)
                fid  = _songs_meta_file_id
            if not fid:
                _, fid = drv.load_songs_meta(svc, CIFRAS_FOLDER_ID)
                with _meta_lock:
                    _songs_meta_file_id = fid
            drv.save_songs_meta(svc, fid, data)
        except Exception:
            pass

    threading.Thread(target=_save, args=(svc,), daemon=True).start()

_songs_meta_loaded = False
_songs_meta_loaded_ts = 0.0
SONGS_META_TTL = 300  # 5 min — re-lê do Drive para sincronizar com outros processos

def _ensure_songs_meta_loaded():
    """Carrega _songs_meta do Drive na primeira chamada, ou quando o TTL expira."""
    global _songs_meta_loaded, _songs_meta_loaded_ts
    now = time.monotonic()
    with _meta_lock:
        if _songs_meta_loaded and (now - _songs_meta_loaded_ts) < SONGS_META_TTL:
            return
        _songs_meta_loaded = True
        _songs_meta_loaded_ts = now
    _load_songs_meta_from_drive()

def _set_song_meta(file_id: str, meta: dict, persist: bool = False, svc=None):
    with _meta_lock:
        _songs_meta[file_id] = {
            "artist":  meta.get("artist", ""),
            "key":     meta.get("key", ""),
            "capo":    meta.get("capo", ""),
            "youtube": meta.get("youtube", ""),
        }
    if persist:
        _persist_songs_meta(svc)

def _get_song_meta(file_id: str) -> dict:
    with _meta_lock:
        return _songs_meta.get(file_id, {})

def _get_library():
    """Retorna a biblioteca escaneada, usando cache quando possível."""
    now = time.monotonic()
    with _cache_lock:
        if _library_cache["data"] is None or (now - _library_cache["ts"]) > CACHE_TTL:
            try:
                if _use_drive():
                    import drive
                    data = drive.scan_library(get_service(), CIFRAS_FOLDER_ID)
                else:
                    data = scan_library_local()
                _library_cache["data"] = data
                _library_cache["ts"] = now
            except Exception as e:
                app.logger.error("Erro ao escanear biblioteca: %s", e)
                # Se já tinha cache antigo, retorna ele em vez de falhar
                if _library_cache["data"] is not None:
                    return _library_cache["data"]
                raise
        return _library_cache["data"]

def invalidate_library_cache():
    with _cache_lock:
        _library_cache["data"] = None
    invalidate_bundle_cache()


# ---------------------------------------------------------------------------
# Cache de bundle (conteúdo completo para sync offline)
# ---------------------------------------------------------------------------

_bundle_cache = {"etag": None, "ts": 0}  # só ETag — não guarda bytes em memória
_bundle_lock  = threading.Lock()

def invalidate_bundle_cache():
    with _bundle_lock:
        _bundle_cache["etag"] = None
        _bundle_cache["ts"]   = 0

def _compute_bundle_etag(songs):
    parts = sorted(
        f"{s['fileId']}:{s.get('modifiedTime', '')}"
        for s in songs if s.get("fileId")
    )
    return hashlib.sha256("\n".join(parts).encode()).hexdigest()[:16]


# ---------------------------------------------------------------------------
# Pastas (criar, renomear, excluir categorias)
# ---------------------------------------------------------------------------

@app.route("/api/folders", methods=["POST"])
@login_required
@owner_required
def api_folder_create():
    """Cria nova categoria dentro de uma seção."""
    data = request.get_json(force=True)
    section = (data.get("section") or "").strip()
    name = (data.get("name") or "").strip()
    if not section or not name:
        return jsonify({"error": "section e name obrigatórios"}), 400

    if _use_drive():
        import drive as drv
        svc = get_service()
        section_id = drv.find_folder_by_name(svc, section, CIFRAS_FOLDER_ID)
        if not section_id:
            return jsonify({"error": f"Seção '{section}' não encontrada"}), 404
        if drv.find_folder_by_name(svc, name, section_id):
            return jsonify({"error": "Categoria já existe"}), 409
        drv.create_folder(svc, name, section_id)
    else:
        folder_path = Path(CIFRAS_ROOT) / section / name
        if not (Path(CIFRAS_ROOT) / section).exists():
            return jsonify({"error": f"Seção '{section}' não encontrada"}), 404
        if folder_path.exists():
            return jsonify({"error": "Categoria já existe"}), 409
        folder_path.mkdir(parents=False)

    invalidate_library_cache()
    return jsonify({"ok": True, "section": section, "name": name}), 201


@app.route("/api/folders/<path:section>/<path:category>", methods=["PUT"])
@login_required
@owner_required
def api_folder_rename(section, category):
    """Renomeia uma categoria."""
    data = request.get_json(force=True)
    new_name = (data.get("newName") or "").strip()
    if not new_name:
        return jsonify({"error": "newName obrigatório"}), 400

    if _use_drive():
        import drive as drv
        svc = get_service()
        section_id = drv.find_folder_by_name(svc, section, CIFRAS_FOLDER_ID)
        if not section_id:
            return jsonify({"error": "Seção não encontrada"}), 404
        folder_id = drv.find_folder_by_name(svc, category, section_id)
        if not folder_id:
            return jsonify({"error": "Categoria não encontrada"}), 404
        if drv.find_folder_by_name(svc, new_name, section_id):
            return jsonify({"error": "Já existe categoria com esse nome"}), 409
        drv.rename_folder(svc, folder_id, new_name)
    else:
        old_path = Path(CIFRAS_ROOT) / section / category
        new_path = Path(CIFRAS_ROOT) / section / new_name
        if not old_path.exists():
            return jsonify({"error": "Categoria não encontrada"}), 404
        if new_path.exists():
            return jsonify({"error": "Já existe categoria com esse nome"}), 409
        old_path.rename(new_path)

    invalidate_library_cache()
    return jsonify({"ok": True, "newName": new_name})


@app.route("/api/folders/<path:section>/<path:category>", methods=["DELETE"])
@login_required
@owner_required
def api_folder_delete(section, category):
    """Exclui uma categoria — apenas se estiver vazia."""
    if _use_drive():
        import drive as drv
        svc = get_service()
        section_id = drv.find_folder_by_name(svc, section, CIFRAS_FOLDER_ID)
        if not section_id:
            return jsonify({"error": "Seção não encontrada"}), 404
        folder_id = drv.find_folder_by_name(svc, category, section_id)
        if not folder_id:
            return jsonify({"error": "Categoria não encontrada"}), 404
        if not drv.is_folder_empty(svc, folder_id):
            return jsonify({"error": "A pasta não está vazia"}), 409
        drv.delete_folder(svc, folder_id)
    else:
        folder_path = Path(CIFRAS_ROOT) / section / category
        if not folder_path.exists():
            return jsonify({"error": "Categoria não encontrada"}), 404
        if any(folder_path.iterdir()):
            return jsonify({"error": "A pasta não está vazia"}), 409
        folder_path.rmdir()

    invalidate_library_cache()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Operações em arquivos de cifra (renomear, excluir, copiar, mover)
# ---------------------------------------------------------------------------

@app.route("/api/songs/delete", methods=["POST"])
@login_required
@owner_required
def api_song_delete():
    data = request.get_json(force=True)
    shortcut_file_id = (data.get("shortcutFileId") or "").strip()
    file_id          = (data.get("fileId") or "").strip()
    path             = (data.get("path") or "").strip()
    if shortcut_file_id:
        # Excluir atalho: apenas o atalho é enviado para a lixeira
        import drive as drv
        drv.trash_file(get_service(), shortcut_file_id)
    elif file_id:
        # Excluir original: manda para lixeira + todos os atalhos que apontam para ele
        import drive as drv
        svc = get_service()
        drv.trash_file(svc, file_id)
        for sc_id in drv.find_shortcuts_to(svc, file_id):
            drv.trash_file(svc, sc_id)
    elif path:
        if not is_safe_path(path):
            return jsonify({"error": "Caminho não permitido"}), 403
        p = Path(path)
        if not p.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404
        p.unlink()
    else:
        return jsonify({"error": "fileId ou path obrigatório"}), 400
    invalidate_library_cache()
    return jsonify({"ok": True})


@app.route("/api/songs/rename", methods=["POST"])
@login_required
@owner_required
def api_song_rename():
    data = request.get_json(force=True)
    shortcut_file_id = (data.get("shortcutFileId") or "").strip()
    file_id          = (data.get("fileId") or "").strip()
    path             = (data.get("path") or "").strip()
    new_name = (data.get("newName") or "").strip()
    if not new_name:
        return jsonify({"error": "newName obrigatório"}), 400
    if file_id:
        import drive as drv
        svc = get_service()
        # Renomeia sempre o arquivo original (targetId)
        current = drv.get_file_name(svc, file_id)
        ext = Path(current).suffix or ".md"
        drv.rename_file(svc, file_id, new_name + ext)
        # Se vier de um atalho, sincroniza o nome do atalho também
        if shortcut_file_id:
            drv.rename_file(svc, shortcut_file_id, new_name + ext)
    elif path:
        if not is_safe_path(path):
            return jsonify({"error": "Caminho não permitido"}), 403
        p = Path(path)
        if not p.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404
        new_path = p.parent / (new_name + p.suffix)
        if new_path.exists():
            return jsonify({"error": "Já existe arquivo com esse nome"}), 409
        p.rename(new_path)
    else:
        return jsonify({"error": "fileId ou path obrigatório"}), 400
    invalidate_library_cache()
    return jsonify({"ok": True})


@app.route("/api/songs/copy", methods=["POST"])
@login_required
@owner_required
def api_song_copy():
    import shutil
    data = request.get_json(force=True)
    file_id = (data.get("fileId") or "").strip()
    path = (data.get("path") or "").strip()
    target_section = (data.get("targetSection") or "").strip()
    target_category = (data.get("targetCategory") or "").strip()
    if not target_section:
        return jsonify({"error": "targetSection obrigatório"}), 400
    if file_id:
        import drive as drv
        svc = get_service()
        target_id = drv.resolve_folder(svc, target_section, target_category or None, CIFRAS_FOLDER_ID)
        fname = drv.get_file_name(svc, file_id)
        drv.create_shortcut(svc, fname, file_id, target_id)
    elif path:
        if not is_safe_path(path):
            return jsonify({"error": "Caminho não permitido"}), 403
        src = Path(path)
        if not src.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404
        dest_dir = Path(CIFRAS_ROOT) / target_section / target_category if target_category else Path(CIFRAS_ROOT) / target_section
        if not dest_dir.exists():
            return jsonify({"error": "Pasta destino não encontrada"}), 404
        dest = dest_dir / src.name
        if dest.exists():
            return jsonify({"error": "Já existe arquivo com esse nome no destino"}), 409
        shutil.copy2(str(src), str(dest))
    else:
        return jsonify({"error": "fileId ou path obrigatório"}), 400
    invalidate_library_cache()
    return jsonify({"ok": True})


@app.route("/api/songs/move", methods=["POST"])
@login_required
@owner_required
def api_song_move():
    import shutil
    data = request.get_json(force=True)
    # Para atalhos, move o atalho em si (não o arquivo original)
    file_id = (data.get("shortcutFileId") or data.get("fileId") or "").strip()
    path = (data.get("path") or "").strip()
    source_section = (data.get("sourceSection") or "").strip()
    source_category = (data.get("sourceCategory") or "").strip()
    target_section = (data.get("targetSection") or "").strip()
    target_category = (data.get("targetCategory") or "").strip()
    if not target_section:
        return jsonify({"error": "targetSection obrigatório"}), 400
    if file_id:
        import drive as drv
        svc = get_service()
        source_id = drv.resolve_folder(svc, source_section, source_category or None, CIFRAS_FOLDER_ID)
        target_id = drv.resolve_folder(svc, target_section, target_category or None, CIFRAS_FOLDER_ID)
        drv.move_file(svc, file_id, source_id, target_id)
    elif path:
        if not is_safe_path(path):
            return jsonify({"error": "Caminho não permitido"}), 403
        src = Path(path)
        if not src.exists():
            return jsonify({"error": "Arquivo não encontrado"}), 404
        dest_dir = Path(CIFRAS_ROOT) / target_section / target_category if target_category else Path(CIFRAS_ROOT) / target_section
        if not dest_dir.exists():
            return jsonify({"error": "Pasta destino não encontrada"}), 404
        dest = dest_dir / src.name
        if dest.exists():
            return jsonify({"error": "Já existe arquivo com esse nome no destino"}), 409
        shutil.move(str(src), str(dest))
    else:
        return jsonify({"error": "fileId ou path obrigatório"}), 400
    invalidate_library_cache()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Segurança
# ---------------------------------------------------------------------------

def is_safe_path(path_str):
    try:
        return Path(path_str).resolve().is_relative_to(Path(CIFRAS_ROOT).resolve())
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _esc(s):
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


_CHORD_WORD_RE = re.compile(r'^[A-G][b#]?(?:m|maj|min|dim|aug|sus|add|[0-9]|/[A-G][b#]?)*$')

def _is_chord_line(line: str) -> bool:
    words = line.strip().split()
    if not words:
        return False
    matches = sum(1 for w in words if _CHORD_WORD_RE.match(re.sub(r'[()[\]]', '', w)))
    return matches / len(words) >= 0.5


def _render_cifra_html(text: str) -> str:
    """Converte texto de cifra em HTML com linhas de acordes coloridas."""
    lines = []
    for line in text.split("\n"):
        if _is_chord_line(line):
            lines.append('<span class="chord-line">' + _esc(line) + '</span>')
        else:
            lines.append(_esc(line))
    return "\n".join(lines)


def _mime_to_ext(mime):
    return {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "application/pdf": ".pdf",
        "text/plain": ".txt",
        "text/markdown": ".md",
    }.get(mime, ".txt")


# ---------------------------------------------------------------------------
# Keep-alive — evita spin-down do Render free tier (idle após 15 min)
# ---------------------------------------------------------------------------

@app.route("/ping")
def ping():
    return "ok", 200


def _start_keep_alive():
    external_url = os.environ.get("EXTERNAL_URL", "").rstrip("/")
    if not external_url:
        return

    import requests as _req

    def _loop():
        while True:
            time.sleep(14 * 60)
            try:
                _req.get(external_url + "/ping", timeout=10)
            except Exception:
                pass

    threading.Thread(target=_loop, daemon=True).start()


_start_keep_alive()


# ---------------------------------------------------------------------------
# Rotas principais
# ---------------------------------------------------------------------------

@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/google<token>.html")
def google_site_verification(token):
    """Verificação de propriedade do Google Search Console (HTML file method)."""
    return Response(
        f"google-site-verification: google{token}.html",
        mimetype="text/html"
    )


@app.route("/terms")
def terms():
    return render_template("terms.html")


@app.route("/")
def index():
    cal_kw_raw = os.environ.get("CALENDAR_KEYWORDS", "").strip()
    cal_keywords = cal_kw_raw if cal_kw_raw else ""
    if not is_oauth_configured():
        return render_template("index.html", user={}, is_owner=True, cal_keywords=cal_keywords)
    if session.get("token"):
        user = current_user()
        log.info("[index] sessão encontrada para %s", user.get("email"))
        return render_template("index.html", user=user, is_owner=is_owner(), cal_keywords=cal_keywords)
    log.info("[index] sessão não encontrada — exibindo landing (cookie=%s)",
             bool(request.cookies.get("session")))
    return render_template(
        "landing.html",
        google_verification=os.environ.get("GOOGLE_SITE_VERIFICATION", "")
    )


@app.route("/api/library")
@login_required
def api_library():
    try:
        return jsonify(_get_library())
    except Exception as e:
        err = str(e)
        cause = str(e.__cause__) if getattr(e, "__cause__", None) else ""
        full = (err + " " + cause).lower()
        app.logger.error("api_library falhou: %s | causa: %s", err, cause)
        _AUTH_SIGNALS = ("invalid_grant", "token has been expired", "sessão expirada",
                         "unauthorized", "401", "revoked", "invalid_token")
        if any(s in full for s in _AUTH_SIGNALS):
            session.clear()
            return jsonify({"error": "Sessão expirada", "login_url": "/"}), 401
        return jsonify({"error": f"Erro ao carregar biblioteca: {err}"}), 500


@app.route("/api/songs")
@login_required
def api_songs():
    _ensure_songs_meta_loaded()
    views = _load_views()
    songs = []
    for s in flatten_songs(_get_library()):
        fid = s.get("fileId", "")
        # Arquivos locais .md: lê frontmatter direto do disco
        if not s.get("artist") and not s.get("key") and s.get("path","").endswith(".md"):
            meta = _md_meta(s["path"])
            s["artist"] = meta.get("artist", "")
            s["key"] = meta.get("key", "")
            s["youtube"] = meta.get("youtube", "")
            if meta.get("title"):
                s["name"] = meta["title"]
        # Arquivos do Drive: injeta metadados do cache global (populado por update_meta e bundle)
        if fid:
            cached = _get_song_meta(fid)
            if cached:
                s.setdefault("artist",  cached.get("artist", ""))
                s.setdefault("key",     cached.get("key", ""))
                s.setdefault("capo",    cached.get("capo", ""))
                s.setdefault("youtube", cached.get("youtube", ""))
                if cached.get("artist"):  s["artist"]  = cached["artist"]
                if cached.get("key"):     s["key"]     = cached["key"]
                if cached.get("capo"):    s["capo"]    = cached["capo"]
                if cached.get("youtube"): s["youtube"] = cached["youtube"]
        s["views"] = views.get(_song_key(s), 0)
        songs.append(s)
    return jsonify(songs)


@app.route("/api/track_view", methods=["POST"])
@login_required
def api_track_view():
    data = request.get_json(force=True)
    key = _song_key(data)
    if not key:
        return jsonify({"error": "fileId ou path obrigatório"}), 400
    new_count = _increment_view(key)
    return jsonify({"views": new_count})


@app.route("/api/save_content", methods=["POST"])
@login_required
@owner_required
def api_save_content():
    """Salva o conteúdo editado (com ou sem transposição), preservando frontmatter."""
    data = request.get_json(force=True)
    new_body = data.get("text", "").strip()
    file_id = data.get("fileId", "").strip()
    path = data.get("path", "").strip()

    if not new_body:
        return jsonify({"error": "Texto vazio"}), 400

    if file_id:
        import drive
        svc = get_service()
        try:
            raw = drive.download_bytes(svc, file_id).decode("utf-8", errors="replace")
        except Exception:
            raw = ""
        if raw.startswith("---"):
            parts = raw.split("---", 2)
            frontmatter = "---" + parts[1] + "---\n\n" if len(parts) >= 3 else ""
        else:
            frontmatter = ""
        try:
            drive.update_md_content(svc, file_id, frontmatter + new_body)
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        invalidate_bundle_cache()
        return jsonify({"ok": True})

    if path:
        try:
            abs_path = Path(path).resolve()
            root = Path(CIFRAS_ROOT).resolve()
            if root not in abs_path.parents and abs_path != root:
                return jsonify({"error": "Caminho não permitido"}), 403
        except Exception:
            return jsonify({"error": "Caminho inválido"}), 400
        if not abs_path.exists() or abs_path.suffix.lower() != ".md":
            return jsonify({"error": "Arquivo não encontrado ou não é .md"}), 404
        original = abs_path.read_text(encoding="utf-8")
        if original.startswith("---"):
            parts = original.split("---", 2)
            frontmatter = "---" + parts[1] + "---\n\n" if len(parts) >= 3 else ""
        else:
            frontmatter = ""
        abs_path.write_text(frontmatter + new_body, encoding="utf-8")
        invalidate_bundle_cache()
        return jsonify({"ok": True})

    return jsonify({"error": "Informe fileId ou path"}), 400


@app.route("/api/cifra")
@login_required
def api_cifra():
    file_id = request.args.get("fileId", "")
    path = request.args.get("path", "")

    if file_id:
        import drive
        mime = request.args.get("mimeType", "")
        try:
            if mime == drive.GDOCS_MIME:
                text = drive.export_gdoc_as_text(get_service(), file_id)
            else:
                content_bytes = drive.download_bytes(get_service(), file_id)
                ext = _mime_to_ext(mime)
                text = extract_text_from_bytes(content_bytes, ext)
        except Exception as e:
            if _is_auth_error(e):
                return _auth_error_response()
            return jsonify({"error": str(e)}), 500
        # Para arquivos .md do Drive, extrai frontmatter igual aos arquivos locais
        is_md = mime in ("text/markdown", "text/plain") or (text.startswith("---") and "\n---" in text)
        if is_md and text.startswith("---"):
            body, meta = _parse_frontmatter(text)
            return jsonify({"text": body, "artist": meta.get("artist", ""), "key": meta.get("key", ""),
                            "title": meta.get("title", ""), "tags": meta.get("tags", []), "capo": meta.get("capo", ""),
                            "youtube": meta.get("youtube", "")})
        return jsonify({"text": text})

    if not path:
        return jsonify({"error": "path ou fileId é obrigatório"}), 400
    if not is_safe_path(path):
        return jsonify({"error": "Acesso negado"}), 403
    if not Path(path).exists():
        return jsonify({"error": "Arquivo não encontrado"}), 404

    p = Path(path)
    if p.suffix.lower() == ".md":
        raw = p.read_text(encoding="utf-8", errors="replace")
        body, meta = _parse_frontmatter(raw)
        return jsonify({"text": body, "artist": meta.get("artist",""), "key": meta.get("key",""),
                        "title": meta.get("title",""), "tags": meta.get("tags",[]), "capo": meta.get("capo",""),
                        "youtube": meta.get("youtube","")})
    return jsonify({"text": extract_text(path), "artist": "", "key": "", "title": "", "tags": [], "youtube": ""})


@app.route("/api/cifras/bundle")
@login_required
def api_cifras_bundle():
    """Retorna todas as cifras em um único JSON para sync offline.
    Suporta If-None-Match / ETag para retornar 304 quando não houve alterações.
    """
    if not _use_drive():
        return jsonify({"error": "Drive não configurado"}), 404

    import drive as drv
    from googleapiclient.discovery import build as _gdrive_build
    from auth import get_credentials as _get_oauth_creds
    from google.auth.transport.requests import Request as _GRequest

    songs = [s for s in flatten_songs(_get_library()) if s.get("fileId")]
    etag  = _compute_bundle_etag(songs)

    # 304 se cliente já tem esta versão
    if request.headers.get("If-None-Match", "") == etag:
        return Response(status=304)

    # Obtém credenciais no contexto Flask antes de spawnar threads
    creds = _get_oauth_creds()
    if creds and creds.expired and creds.refresh_token:
        creds.refresh(_GRequest())

    def _fetch(song):
        local_svc = _gdrive_build("drive", "v3", credentials=creds, cache_discovery=False)
        fid  = song["fileId"]
        mime = song.get("mimeType", "")
        try:
            if mime == drv.GDOCS_MIME:
                text = drv.export_gdoc_as_text(local_svc, fid)
                body, meta = text, {}
            else:
                raw_bytes = drv.download_bytes(local_svc, fid)
                ext = _mime_to_ext(mime)
                if ext in (".md", ".txt"):
                    raw = raw_bytes.decode("utf-8", errors="replace")
                    if raw.startswith("---") and "\n---" in raw:
                        body, meta = _parse_frontmatter(raw)
                    else:
                        body, meta = raw, {}
                else:
                    body = extract_text_from_bytes(raw_bytes, ext)
                    meta = {}
            return fid, {
                "text":    body,
                "name":    song["name"],
                "artist":  meta.get("artist", ""),
                "key":     meta.get("key", ""),
                "capo":    meta.get("capo", ""),
                "youtube": meta.get("youtube", ""),
            }
        except Exception:
            return fid, None

    bundle = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as pool:
        for fid, data in pool.map(_fetch, songs):
            if data:
                bundle[fid] = data
                _set_song_meta(fid, data)  # alimenta cache em memória

    _persist_songs_meta(svc=_gdrive_build("drive", "v3", credentials=creds, cache_discovery=False))  # salva metadados atualizados no Drive em background

    with _bundle_lock:
        _bundle_cache["etag"] = etag
        _bundle_cache["ts"]   = time.monotonic()

    json_bytes = json.dumps({"etag": etag, "songs": bundle}, ensure_ascii=False).encode("utf-8")
    resp = Response(json_bytes, content_type="application/json")
    resp.headers["ETag"] = etag
    resp.headers["Cache-Control"] = "no-store"
    return resp


@app.route("/api/upload", methods=["POST"])
@login_required
@owner_required
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400
    file = request.files["file"]
    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return jsonify({"error": "Formato não suportado"}), 400
    text = extract_text_from_bytes(file.read(), suffix)
    return jsonify({"text": text, "name": Path(file.filename).stem})


def _build_export_html(songs, title, auto_print=False):
    """Gera o HTML estilizado do repertório (usado pelo endpoint HTML e PDF)."""
    today = date.today().strftime("%d/%m/%Y")

    # Inline logo SVG
    logo_path = Path(__file__).parent / "static" / "brand" / "logo-mono-dark.svg"
    try:
        logo_svg = logo_path.read_text(encoding="utf-8")
        # Remove XML comments to keep it clean inline
        import re as _re
        logo_svg = _re.sub(r'<!--.*?-->', '', logo_svg, flags=_re.DOTALL).strip()
    except Exception:
        logo_svg = ""

    def _song_card(idx, s):
        num     = idx + 1
        name    = _esc(s.get("name", ""))
        note    = _esc((s.get("note") or "").strip())
        key     = s.get("key", "")
        capo    = int(s.get("capo") or 0)
        meta_parts = []
        if note:
            meta_parts.append(f'<span class="badge badge-note">{note}</span>')
        if key:
            meta_parts.append(f'<span class="badge badge-key">Tom: {_esc(key)}</span>')
        if capo > 0:
            meta_parts.append(f'<span class="badge badge-capo">Capotraste na {capo}ª casa</span>')
        meta_row = f'<div class="song-meta">{" ".join(meta_parts)}</div>' if meta_parts else ""
        cifra_html = _render_cifra_html(s.get("text", ""))
        return (
            f'<div class="song">'
            f'  <div class="song-header">'
            f'    <h2><span class="song-num">{num}.</span> {name}</h2>'
            f'  </div>'
            f'  {meta_row}'
            f'  <pre>{cifra_html}</pre>'
            f'</div>\n'
        )

    n = len(songs)
    song_word = "música" if n == 1 else "músicas"

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_esc(title)}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Roboto+Mono:wght@400;700;800&display=swap');

  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

  body {{
    font-family: 'Inter', system-ui, sans-serif;
    background: #f4f2fa;
    color: #1a1d2e;
    max-width: 860px;
    margin: 0 auto;
    padding: 0 0 80px;
  }}

  /* ── Cabeçalho com fundo roxo ── */
  .doc-banner {{
    background: #5b4b8a;
    padding: 28px 40px 24px;
    display: flex;
    align-items: flex-end;
    justify-content: space-between;
    gap: 16px;
  }}
  .doc-banner .logo svg {{ height: 56px; width: auto; filter: brightness(10); opacity: .92; }}
  .doc-banner-right {{ text-align: right; }}
  .doc-title {{
    font-size: 1.6em;
    font-weight: 800;
    letter-spacing: -.03em;
    color: #fff;
    line-height: 1.15;
  }}
  .doc-subtitle {{
    font-size: .8em;
    color: rgba(255,255,255,.65);
    margin-top: 4px;
  }}

  .doc-body {{ padding: 32px 40px 0; }}

  /* ── Song card ── */
  .song {{
    background: #fff;
    border: 1px solid #e4e0f4;
    border-radius: 12px;
    padding: 20px 24px 24px;
    margin-bottom: 24px;
    page-break-inside: avoid;
    break-inside: avoid;
  }}
  .song-header {{
    margin-bottom: 8px;
    padding-bottom: 10px;
    border-bottom: 1px solid #ede9fa;
  }}
  .song h2 {{
    font-size: 1.05em;
    font-weight: 700;
    color: #5b4b8a;
    letter-spacing: -.015em;
    display: flex;
    align-items: baseline;
    gap: 6px;
  }}
  .song-num {{
    color: #d4af37;
    font-size: .9em;
    font-weight: 800;
    flex-shrink: 0;
  }}
  .song-meta {{
    display: flex;
    align-items: center;
    gap: 8px;
    flex-wrap: wrap;
    margin-bottom: 14px;
  }}
  .meta-artist {{
    font-size: .78em;
    font-style: italic;
    color: #7a6fa8;
  }}
  .badge {{
    font-size: .72em;
    font-weight: 600;
    padding: 2px 9px;
    border-radius: 99px;
    white-space: nowrap;
  }}
  .badge-cat {{
    background: rgba(91,75,138,.1);
    color: #5b4b8a;
    border: 1px solid rgba(91,75,138,.2);
  }}
  .badge-key {{
    background: rgba(212,175,55,.12);
    color: #9a7a10;
    border: 1px solid rgba(212,175,55,.3);
    font-weight: 700;
  }}
  .badge-capo {{
    background: #d4af37;
    color: #fff;
    border: 1px solid #b8942a;
    font-weight: 700;
  }}
  .badge-note {{
    background: rgba(91,75,138,.08);
    color: #5b4b8a;
    border: 1px solid rgba(91,75,138,.18);
    font-style: italic;
  }}

  /* ── Cifra ── */
  pre {{
    font-family: 'Roboto Mono', 'Consolas', 'Courier New', monospace;
    font-size: .8em;
    line-height: 1.35;
    white-space: pre-wrap;
    word-break: break-word;
    color: #2e2645;
    font-weight: 400;
  }}
  .chord-line {{ color: #5b4b8a; font-weight: 800; }}

  /* ── Footer ── */
  .doc-footer {{
    margin: 40px 40px 0;
    padding-top: 14px;
    border-top: 1px solid #e2dff5;
    font-size: .73em;
    color: #b0a8cc;
    text-align: center;
  }}

  /* ── Botão voltar (mobile) ── */
  .btn-back {{
    display: none;
    position: fixed; bottom: 24px; right: 20px;
    background: #5b4b8a; color: #fff;
    border: none; border-radius: 50px;
    padding: .65rem 1.2rem;
    font-family: 'Inter', system-ui, sans-serif;
    font-size: .85rem; font-weight: 700;
    box-shadow: 0 4px 16px rgba(91,75,138,.4);
    cursor: pointer; z-index: 999;
    text-decoration: none;
    touch-action: manipulation;
  }}
  @media (max-width: 768px) {{
    .doc-banner {{ padding: 20px; }}
    .doc-body {{ padding: 20px 16px 0; }}
    .doc-footer {{ margin: 32px 16px 0; }}
    .btn-back {{ display: inline-flex; align-items: center; gap: .4rem; }}
  }}

  /* ── Print / WeasyPrint ── */
  @page {{
    size: A4 portrait;
    margin: 1.1cm 1.4cm 1.1cm 1.4cm;
  }}
  @media print {{
    body {{
      background: #fff;
      max-width: 100%;
      padding: 0;
    }}
    .doc-banner {{
      padding: 12px 18px 10px;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
    }}
    .doc-banner .logo svg {{ height: 36px; }}
    .doc-title {{ font-size: 1.25em; }}
    .doc-body {{ padding: 10px 0 0; }}
    .song {{
      border: 1px solid #ddd;
      margin-bottom: 10px;
      padding: 10px 14px 12px;
      border-radius: 8px;
    }}
    .song-meta {{ margin-bottom: 8px; }}
    .doc-footer {{ margin: 16px 0 0; }}
    .btn-back {{ display: none !important; }}
    .badge {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  }}
</style>
</head>
<body>

<header class="doc-banner">
  <div>
    <div class="logo">{logo_svg}</div>
  </div>
  <div class="doc-banner-right">
    <div class="doc-title">{_esc(title)}</div>
    <div class="doc-subtitle">{n} {song_word} · {_esc(today)}</div>
  </div>
</header>

<div class="doc-body">
{"".join(_song_card(i, s) for i, s in enumerate(songs))}
</div>

<footer class="doc-footer">My Cifras · gerado em {_esc(today)}</footer>

<a class="btn-back" onclick="window.close(); history.back(); return false;" href="#">← Voltar</a>

{'<script>window.onload=function(){window.print();}</script>' if auto_print else ''}
</body>
</html>"""

    return html


@app.route("/api/export", methods=["POST"])
@login_required
def api_export():
    data = request.get_json(force=True)
    songs = data.get("songs", [])
    title = data.get("title", "Repertório")
    auto_print = data.get("print", False)
    html = _build_export_html(songs, title, auto_print)
    return Response(html, mimetype="text/html; charset=utf-8")


@app.route("/api/export/pdf", methods=["POST"])
@login_required
def api_export_pdf():
    """Gera PDF diretamente via WeasyPrint — sem precisar abrir o browser."""
    import io
    import re as _re
    from flask import send_file
    data  = request.get_json(force=True)
    songs = data.get("songs", [])
    title = data.get("title", "Repertório")
    html  = _build_export_html(songs, title, auto_print=False)
    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(
            string=html,
            base_url=request.url_root,
        ).write_pdf()
    except Exception as e:
        app.logger.error("WeasyPrint indisponível (%s) — retornando HTML como fallback", e)
        # Fallback: devolve HTML para o cliente abrir em nova aba
        return Response(
            html,
            mimetype="text/html; charset=utf-8",
            headers={"X-Export-Fallback": "1"},
        )
    buf = io.BytesIO(pdf_bytes)
    buf.seek(0)
    safe_name = _re.sub(r'[<>:"/\\|?*]', "_", title).strip() or "repertorio"
    return send_file(
        buf,
        as_attachment=True,
        download_name=safe_name + ".pdf",
        mimetype="application/pdf",
    )


# ---------------------------------------------------------------------------
# Rotas de importação
# ---------------------------------------------------------------------------

@app.route("/api/import/preview", methods=["POST"])
@login_required
def api_import_preview():
    data = request.get_json(force=True)
    url = data.get("url", "").strip()
    text = data.get("text", "").strip()

    if url:
        try:
            from scraper import fetch_cifra
            return jsonify(fetch_cifra(url))
        except Exception as e:
            return jsonify({"error": f"Não foi possível buscar a URL: {e}"}), 400

    if text:
        return jsonify({"title": "", "artist": "", "key": "", "text": text})

    return jsonify({"error": "Envie url ou text"}), 400


@app.route("/api/import/save", methods=["POST"])
@login_required
@owner_required
def api_import_save():
    data = request.get_json(force=True)
    title    = data.get("title", "").strip()
    artist   = data.get("artist", "").strip()
    key      = data.get("key", "").strip()
    capo     = str(data.get("capo", "")).strip()
    youtube  = data.get("youtube", "").strip()
    section  = data.get("section", "").strip()
    category = data.get("category", "").strip()
    text     = data.get("text", "").strip()

    if not title:
        return jsonify({"error": "Título é obrigatório"}), 400
    if not section:
        return jsonify({"error": "Seção é obrigatória"}), 400

    cat_display = category if category and category != "_raiz" else ""
    content = (
        "---\n"
        f"title: {title}\n"
        f"artist: {artist}\n"
        f"key: {key}\n"
        f"capo: {capo}\n"
        f"youtube: {youtube}\n"
        f"section: {section}\n"
        f"category: {cat_display}\n"
        "---\n\n"
        + text
    )

    meta = {"artist": artist, "key": key, "capo": capo, "youtube": youtube}

    if _use_drive():
        import drive
        svc = get_service()
        folder_id = drive.resolve_folder(svc, section, category or "_raiz", CIFRAS_FOLDER_ID)
        file_id = drive.upload_md(svc, title, content, folder_id)
        _ensure_songs_meta_loaded()
        _set_song_meta(file_id, meta, persist=True, svc=svc)
        invalidate_library_cache()
        return jsonify({"ok": True, "fileId": file_id})
    else:
        safe_name = re.sub(r'[<>:"/\\|?*]', "_", title).strip()
        dest_dir = Path(CIFRAS_ROOT) / section / (category if category and category != "_raiz" else "")
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_file = dest_dir / (safe_name + ".md")
        dest_file.write_text(content, encoding="utf-8")
        invalidate_library_cache()
        return jsonify({"ok": True, "path": str(dest_file)})


@app.route("/api/export/docx", methods=["POST"])
@login_required
def api_export_docx():
    """Gera .docx minimalista — fácil de editar no Word."""
    import io
    import re as _re
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from flask import send_file

    PURPLE = RGBColor(0x5b, 0x4b, 0x8a)
    DARK   = RGBColor(0x1a, 0x1d, 0x2e)

    _CHORD_TOKEN = _re.compile(
        r'^[A-G][b#]?(?:m|maj|min|dim|aug|sus|add|[0-9]|/[A-G][b#]?)*$'
    )
    def _is_chord_line(line):
        tokens = line.strip().split()
        return bool(tokens) and all(_CHORD_TOKEN.match(t) for t in tokens)

    def _tight(p, before=0, after=0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        return p

    def _bottom_rule(p, color="c8c0e0", sz=4):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), str(sz))
        bt.set(qn("w:space"), "1"); bt.set(qn("w:color"), color)
        pBdr.append(bt); pPr.append(pBdr)

    data  = request.get_json(force=True)
    songs = data.get("songs", [])
    title = data.get("title", "Repertório")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # ── Título do documento ──────────────────────────────────────────────────
    tp = doc.add_paragraph()
    _tight(tp, before=0, after=6)
    tr = tp.add_run(title)
    tr.font.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = PURPLE
    _bottom_rule(tp)

    # ── Músicas (página por música) ──────────────────────────────────────────
    for i, song in enumerate(songs):
        if i > 0:
            doc.add_page_break()

        # Nome da música
        np_ = doc.add_paragraph()
        _tight(np_, before=0, after=4)
        nr = np_.add_run(f"{i + 1}.  {song.get('name', '')}")
        nr.font.bold = True; nr.font.size = Pt(13); nr.font.color.rgb = PURPLE
        _bottom_rule(np_)

        # Metadados numa linha só (nota | Tom: X | Capotraste na Xª)
        note = (song.get("note") or "").strip()
        key  = (song.get("key")  or "").strip()
        capo = int(song.get("capo") or 0)
        meta_parts = []
        if note: meta_parts.append(note)
        if key:  meta_parts.append(f"Tom: {key}")
        if capo > 0: meta_parts.append(f"Capotraste na {capo}ª casa")
        if meta_parts:
            mp = doc.add_paragraph()
            _tight(mp, before=2, after=6)
            mr = mp.add_run("  " + "  ·  ".join(meta_parts))
            mr.font.size = Pt(9); mr.font.italic = True
            mr.font.color.rgb = RGBColor(0x7a, 0x6f, 0xa8)

        # Cifra
        text = (song.get("text") or "").strip()
        for line in text.split("\n"):
            lp = doc.add_paragraph()
            _tight(lp, before=0, after=0)
            if not line.strip():
                _tight(lp, before=0, after=3)
                continue
            run = lp.add_run(line)
            run.font.name = "Consolas"; run.font.size = Pt(9)
            if _is_chord_line(line):
                run.font.bold = True; run.font.color.rgb = PURPLE
            else:
                run.font.color.rgb = DARK

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _re.sub(r'[<>:"/\\|?*]', "_", title).strip() or "repertorio"
    return send_file(
        buf,
        as_attachment=True,
        download_name=safe_name + ".docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/songs/update_meta", methods=["POST"])
@login_required
@owner_required
def api_update_meta():
    """Atualiza apenas o frontmatter YAML de um arquivo .md."""
    data = request.get_json(force=True)
    file_id = (data.get("fileId") or "").strip()
    path    = (data.get("path") or "").strip()
    new_meta = {
        "title":   (data.get("title")   or "").strip(),
        "artist":  (data.get("artist")  or "").strip(),
        "key":     (data.get("key")     or "").strip(),
        "capo":    str(data.get("capo") or "").strip(),
        "tags":    data.get("tags", []),
        "youtube": (data.get("youtube") or "").strip(),
    }

    def _patch_frontmatter(raw, new_meta):
        """Substitui os campos de metadados no frontmatter, preserva o corpo."""
        body = raw
        existing = {"section": "", "category": ""}
        if raw.startswith("---"):
            parts = raw.split("---", 2)
            if len(parts) >= 3:
                for line in parts[1].splitlines():
                    if ":" in line:
                        k, _, v = line.partition(":")
                        existing[k.strip().lower()] = v.strip().strip('"').strip("'")
                body = parts[2].strip()

        tags_str = ", ".join(new_meta["tags"]) if isinstance(new_meta["tags"], list) else new_meta["tags"]
        # Normaliza tags para lista YAML
        tags_list = [t.strip() for t in tags_str.split(",") if t.strip()] if tags_str else []
        tags_yaml = "[" + ", ".join(tags_list) + "]"

        capo_line    = f"capo: {new_meta['capo']}\n" if new_meta.get('capo') else ""
        youtube_line = f"youtube: {new_meta['youtube']}\n" if new_meta.get('youtube') else ""
        fm = (
            "---\n"
            f"title: {new_meta['title']}\n"
            f"artist: {new_meta['artist']}\n"
            f"key: {new_meta['key']}\n"
            + capo_line +
            f"section: {existing.get('section','')}\n"
            f"category: {existing.get('category','')}\n"
            f"tags: {tags_yaml}\n"
            + youtube_line +
            "---\n\n"
        )
        return fm + body

    if file_id:
        import drive
        svc = get_service()
        try:
            raw = drive.download_bytes(svc, file_id).decode("utf-8", errors="replace")
        except Exception as e:
            if _is_auth_error(e):
                return _auth_error_response()
            return jsonify({"error": f"Erro ao ler arquivo: {e}"}), 500
        patched = _patch_frontmatter(raw, new_meta)
        try:
            drive.update_md_content(svc, file_id, patched)
            # Se o título mudou, renomeia o arquivo .md para corresponder
            if new_meta["title"]:
                current_name = drive.get_file_name(svc, file_id)
                if Path(current_name).stem != new_meta["title"]:
                    drive.rename_file(svc, file_id, new_meta["title"] + ".md")
        except Exception as e:
            if _is_auth_error(e):
                return _auth_error_response()
            return jsonify({"error": str(e)}), 500
        _set_song_meta(file_id, new_meta, persist=True, svc=svc)
        invalidate_library_cache()
        return jsonify({"ok": True})

    if path:
        if not is_safe_path(path):
            return jsonify({"error": "Caminho não permitido"}), 403
        p = Path(path)
        if not p.exists() or p.suffix.lower() != ".md":
            return jsonify({"error": "Arquivo não encontrado ou não é .md"}), 404
        raw = p.read_text(encoding="utf-8")
        patched = _patch_frontmatter(raw, new_meta)
        p.write_text(patched, encoding="utf-8")
        invalidate_library_cache()
        return jsonify({"ok": True})

    return jsonify({"error": "fileId ou path obrigatório"}), 400


@app.route("/api/save_tone", methods=["POST"])
@login_required
@owner_required
def api_save_tone():
    data = request.get_json(force=True)
    new_text = data.get("text", "").strip()
    file_id = data.get("fileId", "").strip()
    path = data.get("path", "").strip()

    if not new_text:
        return jsonify({"error": "Texto vazio"}), 400

    if file_id:
        # Drive: baixa o arquivo original para preservar frontmatter
        import drive
        svc = get_service()
        try:
            raw = drive.download_bytes(svc, file_id)
            original = raw.decode("utf-8", errors="replace")
        except Exception:
            original = ""

        # Preserva o bloco frontmatter se existir
        if original.startswith("---"):
            parts = original.split("---", 2)
            frontmatter = "---" + parts[1] + "---\n\n" if len(parts) >= 3 else ""
        else:
            frontmatter = ""

        content = frontmatter + new_text
        try:
            drive.update_md_content(svc, file_id, content)
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        return jsonify({"ok": True})

    elif path:
        # Local: valida path dentro de CIFRAS_ROOT
        try:
            abs_path = Path(path).resolve()
            root = Path(CIFRAS_ROOT).resolve()
            if root not in abs_path.parents and abs_path != root:
                return jsonify({"error": "Caminho não permitido"}), 403
        except Exception:
            return jsonify({"error": "Caminho inválido"}), 400

        if not abs_path.exists() or abs_path.suffix.lower() != ".md":
            return jsonify({"error": "Arquivo não encontrado ou não é .md"}), 404

        original = abs_path.read_text(encoding="utf-8")
        if original.startswith("---"):
            parts = original.split("---", 2)
            frontmatter = "---" + parts[1] + "---\n\n" if len(parts) >= 3 else ""
        else:
            frontmatter = ""

        abs_path.write_text(frontmatter + new_text, encoding="utf-8")
        return jsonify({"ok": True})

    return jsonify({"error": "Informe fileId ou path"}), 400


@app.route("/api/search/content")
@login_required
def api_search_content():
    """Busca full-text nas letras/conteúdo das cifras.
    Drive: usa fullText contains via API.
    Local: lê os arquivos .md e procura no conteúdo.
    Retorna lista de { name, section, category, fileId/path, mimeType?, excerpt }.
    """
    q = request.args.get("q", "").strip()
    if not q or len(q) < 2:
        return jsonify({"error": "Consulta muito curta (mín. 2 caracteres)"}), 400

    results = []

    if _use_drive():
        import drive as drv
        svc = get_service()
        try:
            items = drv.search_content(svc, q, CIFRAS_FOLDER_ID)
            views = _load_views()
            # Enriquece com section/category a partir da biblioteca cacheada
            lib = _get_library()
            # Monta índice fileId → {section, category, name, key}
            fid_index = {}
            for sec, cats in lib.items():
                for cat, songs in cats.items():
                    for s in songs:
                        if s.get("fileId"):
                            fid_index[s["fileId"]] = {
                                "section": sec,
                                "category": cat,
                                "name": s.get("name", ""),
                                "key": s.get("key", ""),
                                "artist": s.get("artist", ""),
                            }
            for item in items:
                fid = item.get("fileId", "")
                meta = fid_index.get(fid, {})
                results.append({
                    "fileId": fid,
                    "name": meta.get("name") or item.get("name", ""),
                    "section": meta.get("section", ""),
                    "category": meta.get("category", ""),
                    "key": meta.get("key", ""),
                    "artist": meta.get("artist", ""),
                    "mimeType": item.get("mimeType", ""),
                    "excerpt": item.get("excerpt", ""),
                    "views": views.get(fid, 0),
                })
        except Exception as e:
            return jsonify({"error": f"Erro na busca: {e}"}), 500
    else:
        # Busca local: percorre os .md dentro de CIFRAS_ROOT
        root = Path(CIFRAS_ROOT)
        if not root.exists():
            return jsonify([])
        ql = _normalize_search(q)
        views = _load_views()
        for section_dir in sorted(root.iterdir()):
            if not section_dir.is_dir():
                continue
            section = section_dir.name
            for item in sorted(section_dir.rglob("*")):
                if not item.is_file() or item.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                rel = item.relative_to(section_dir)
                category = rel.parts[0] if len(rel.parts) > 1 else "_raiz"
                try:
                    raw = item.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    continue
                body, meta = _parse_frontmatter(raw)
                searchable = _normalize_search(body + " " + meta.get("artist", "") + " " + meta.get("title", ""))
                if ql not in searchable:
                    continue
                # Gera excerpt: linha que contém o termo
                excerpt = ""
                for line in body.splitlines():
                    if ql in _normalize_search(line):
                        excerpt = line.strip()[:120]
                        break
                path_str = str(item)
                results.append({
                    "path": path_str,
                    "name": meta.get("title") or item.stem,
                    "section": section,
                    "category": category,
                    "key": meta.get("key", ""),
                    "artist": meta.get("artist", ""),
                    "excerpt": excerpt,
                    "views": views.get(path_str, 0),
                })

    return jsonify(results)


@app.route("/api/sections")
@login_required
def api_sections():
    if _use_drive():
        import drive
        library = drive.scan_library(get_service(), CIFRAS_FOLDER_ID)
    else:
        library = scan_library_local()
    return jsonify({sec: list(cats.keys()) for sec, cats in library.items()})


# ---------------------------------------------------------------------------
# Liturgia diária
# ---------------------------------------------------------------------------

_liturgia_cache = {}   # { "YYYY-MM-DD": { data, liturgia, cor, leituras } }

_COR_MAP = {
    "branco": "#f0ede6",
    "verde":  "#3d7a52",
    "roxo":   "#6b3fa0",
    "violeta":"#6b3fa0",
    "vermelho":"#b92b2b",
    "rosa":   "#d4608a",
    "preto":  "#2c2c2c",
}

def _cor_hex(cor_str):
    return _COR_MAP.get((cor_str or "").lower().strip(), "#9b97b0")


@app.route("/api/liturgia")
@login_required
def api_liturgia():
    from datetime import datetime, date as _date
    import requests as _req

    date_str = request.args.get("date", "")
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d").date() if date_str else _date.today()
    except ValueError:
        return jsonify({"error": "Formato de data inválido. Use YYYY-MM-DD"}), 400

    cache_key = d.isoformat()
    if cache_key in _liturgia_cache:
        return jsonify(_liturgia_cache[cache_key])

    try:
        resp = _req.get(
            "https://liturgia.up.railway.app/v2/",
            params={"dia": d.day, "mes": d.month, "ano": d.year},
            timeout=8,
        )
        resp.raise_for_status()
        raw = resp.json()
    except Exception as e:
        return jsonify({"error": f"Não foi possível carregar a liturgia: {e}"}), 502

    leituras = raw.get("leituras", {})

    def _reading(key):
        items = leituras.get(key) or []
        if not items:
            return None
        item = items[0]
        return {
            "referencia": item.get("referencia", ""),
            "titulo":     item.get("titulo", ""),
            "texto":      item.get("texto", ""),
            "refrao":     item.get("refrao", ""),   # salmo
        }

    cor = raw.get("cor", "")
    tempo = raw.get("tempo", raw.get("periodo", raw.get("season", "")))
    result = {
        "data":          raw.get("data", cache_key),
        "liturgia":      raw.get("liturgia", ""),
        "cor":           cor,
        "corHex":        _cor_hex(cor),
        "tempo":         tempo,
        "primeiraLeitura": _reading("primeiraLeitura"),
        "salmo":           _reading("salmo"),
        "segundaLeitura":  _reading("segundaLeitura"),
        "evangelho":       _reading("evangelho"),
    }

    _liturgia_cache[cache_key] = result
    return jsonify(result)


# ---------------------------------------------------------------------------
# Google Calendar
# ---------------------------------------------------------------------------

CALENDAR_ID = os.environ.get("GOOGLE_CALENDAR_ID", "primary")

# Palavras-chave para filtrar eventos do calendário (case-insensitive, sem acento)
def _calendar_keywords():
    raw = os.environ.get("CALENDAR_KEYWORDS", "").strip()
    if not raw:
        return []
    import unicodedata
    def _norm(s):
        return unicodedata.normalize("NFD", s).encode("ascii", "ignore").decode().lower()
    return [_norm(k.strip()) for k in raw.split(",") if k.strip()]

def _event_matches_keywords(title, keywords):
    """Retorna True se o título contém ao menos uma palavra-chave (ou se não há filtro)."""
    if not keywords:
        return True
    import unicodedata
    def _norm(s):
        return unicodedata.normalize("NFD", s).encode("ascii", "ignore").decode().lower()
    t = _norm(title)
    return any(kw in t for kw in keywords)

def _format_event(item):
    start = item.get("start", {})
    end   = item.get("end", {})
    all_day = "dateTime" not in start
    return {
        "id":          item.get("id", ""),
        "title":       item.get("summary", "(sem título)"),
        "description": item.get("description", ""),
        "location":    item.get("location", ""),
        "start":       start.get("dateTime") or start.get("date") or "",
        "end":         end.get("dateTime")   or end.get("date")   or "",
        "allDay":      all_day,
        "htmlLink":    item.get("htmlLink", ""),
    }

@app.route("/api/calendar")
@login_required
def api_calendar():
    from datetime import datetime, timezone, timedelta
    try:
        from auth import get_calendar_service
        svc = get_calendar_service()

        # Aceita range customizado (FullCalendar envia start/end ao buscar eventos)
        start_param = request.args.get("start")
        end_param   = request.args.get("end")
        if start_param and end_param:
            time_min = start_param if start_param.endswith("Z") else start_param + "T00:00:00Z"
            time_max = end_param   if end_param.endswith("Z")   else end_param   + "T00:00:00Z"
        else:
            dt_now = datetime.now(timezone.utc)
            time_min = dt_now.isoformat()
            time_max = (dt_now + timedelta(days=30)).isoformat()

        result = svc.events().list(
            calendarId=CALENDAR_ID,
            timeMin=time_min,
            timeMax=time_max,
            maxResults=250,
            singleEvents=True,
            orderBy="startTime",
        ).execute()

        keywords = _calendar_keywords()
        events = [
            _format_event(item) for item in result.get("items", [])
            if _event_matches_keywords(item.get("summary", ""), keywords)
        ]
        return jsonify({"events": events})

    except Exception as e:
        app.logger.error("Erro ao carregar calendário: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route("/api/calendar/events", methods=["POST"])
@login_required
def api_calendar_create():
    from auth import get_calendar_service
    data = request.get_json(force=True)
    try:
        svc = get_calendar_service()
        body = {"summary": data.get("title", "").strip()}
        if data.get("description"): body["description"] = data["description"]
        if data.get("location"):    body["location"]    = data["location"]
        if data.get("allDay"):
            body["start"] = {"date": data["start"][:10]}
            body["end"]   = {"date": data["end"][:10]}
        else:
            body["start"] = {"dateTime": data["start"], "timeZone": data.get("timeZone", "America/Sao_Paulo")}
            body["end"]   = {"dateTime": data["end"],   "timeZone": data.get("timeZone", "America/Sao_Paulo")}
        event = svc.events().insert(calendarId=CALENDAR_ID, body=body).execute()
        return jsonify(_format_event(event))
    except Exception as e:
        app.logger.error("Erro ao criar evento: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route("/api/calendar/events/<event_id>", methods=["PUT"])
@login_required
def api_calendar_update(event_id):
    from auth import get_calendar_service
    data = request.get_json(force=True)
    try:
        svc = get_calendar_service()
        body = {"summary": data.get("title", "").strip()}
        if data.get("description"): body["description"] = data["description"]
        if data.get("location"):    body["location"]    = data["location"]
        if data.get("allDay"):
            body["start"] = {"date": data["start"][:10]}
            body["end"]   = {"date": data["end"][:10]}
        else:
            body["start"] = {"dateTime": data["start"], "timeZone": data.get("timeZone", "America/Sao_Paulo")}
            body["end"]   = {"dateTime": data["end"],   "timeZone": data.get("timeZone", "America/Sao_Paulo")}
        event = svc.events().update(calendarId=CALENDAR_ID, eventId=event_id, body=body).execute()
        return jsonify(_format_event(event))
    except Exception as e:
        app.logger.error("Erro ao atualizar evento: %s", e)
        return jsonify({"error": str(e)}), 500


@app.route("/api/calendar/events/<event_id>", methods=["DELETE"])
@login_required
def api_calendar_delete(event_id):
    if not event_id or event_id == "null":
        return jsonify({"error": "ID de evento inválido"}), 400
    from auth import get_calendar_service
    try:
        svc = get_calendar_service()
        svc.events().delete(calendarId=CALENDAR_ID, eventId=event_id).execute()
        return jsonify({"ok": True})
    except Exception as e:
        app.logger.error("Erro ao excluir evento: %s", e)
        return jsonify({"error": str(e)}), 500


# ---------------------------------------------------------------------------
# Batch fix: detecta e escreve campo "key" nos .md sem tonalidade
# ---------------------------------------------------------------------------

_NOTES_PY       = ["C","C#","D","D#","E","F","F#","G","G#","A","A#","B"]
_ENHARMONIC_PY  = {"Db":"C#","Eb":"D#","Gb":"F#","Ab":"G#","Bb":"A#"}
# Suporta notação brasileira: F7M, F7+, Cm7-, dim, aug, sus, etc.
_CHORD_RE_PY    = re.compile(r'^[A-G][b#]?(m|M|maj|min|dim|aug|sus|add|\d|[+\-°ø]|/[A-G][b#]?)*$')
_INTRO_LINE_PY  = re.compile(r'\b(intro|introdução|introd|int\.)\b', re.IGNORECASE)
_SECTION_RE_PY  = re.compile(
    r'^\s*[\[(]?\s*(refrão|refrao|coro|verso|ponte|bridge|chorus|verse|intro\b|introd|final|solo|instrumental)\s*[\])]?\s*:?\s*$',
    re.IGNORECASE
)
_MAJOR_INTERVALS_PY = [0, 2, 4, 5, 7, 9, 11]

def _normalize_note_py(n):
    return _ENHARMONIC_PY.get(n, n)

def _is_chord_line_py(line):
    if _SECTION_RE_PY.match(line.strip()):
        return False
    words = [w.strip("()[]") for w in line.strip().split() if w.strip("()[]")]
    if not words:
        return False
    hits = sum(1 for w in words if _CHORD_RE_PY.match(w))
    return hits / len(words) >= 0.5

def _detect_key_py(text):
    """Detecta tonalidade por análise harmônica: escolhe escala maior com melhor fit."""
    roots = []
    for line in text.splitlines():
        if _INTRO_LINE_PY.search(line):
            continue
        if _is_chord_line_py(line):
            for m in re.finditer(r'[A-G][b#]?', line):
                roots.append(_normalize_note_py(m.group()))
    if not roots:
        return None
    best_key, best_score = None, -1
    for i, root in enumerate(_NOTES_PY):
        scale = {_NOTES_PY[(i + iv) % 12] for iv in _MAJOR_INTERVALS_PY}
        score = sum(1 for r in roots if r in scale)
        if score > best_score:
            best_score, best_key = score, root
    return best_key

def _parse_md(content):
    """Retorna (frontmatter_dict, body). Frontmatter simples key: value."""
    if not content.startswith("---"):
        return {}, content
    parts = content.split("---", 2)
    if len(parts) < 3:
        return {}, content
    fm = {}
    for line in parts[1].splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            fm[k.strip()] = v.strip().strip("\"'")
    body = parts[2].lstrip("\n")
    return fm, body

def _build_md(fm, body):
    """Reconstrói o .md com frontmatter atualizado."""
    lines = ["---"]
    for k, v in fm.items():
        val = str(v) if v is not None else ""
        lines.append(f"{k}: {val}")
    lines.append("---")
    lines.append("")
    return "\n".join(lines) + body


@app.route("/api/admin/fix-keys", methods=["POST"])
@login_required
@owner_required
def api_fix_keys():
    """Percorre todos os .md do acervo e escreve o campo key nos que não têm."""
    from auth import get_service
    from drive import scan_library, download_bytes, update_md_content

    svc = get_service()
    try:
        lib = scan_library(svc, CIFRAS_FOLDER_ID)
    except Exception as e:
        return jsonify({"error": f"Erro ao escanear biblioteca: {e}"}), 500

    fixed = []
    skipped = []
    errors = []

    for section, cats in lib.items():
        for cat, songs in cats.items():
            for song in songs:
                name = song.get("name", "")
                file_id = song.get("id")
                mime = song.get("mimeType", "")

                # Processa apenas .md
                if not (name.lower().endswith(".md") or mime == "text/markdown"):
                    skipped.append(name)
                    continue

                try:
                    raw = download_bytes(svc, file_id)
                    content = raw.decode("utf-8", errors="replace")
                    fm, body = _parse_md(content)

                    # Já tem key preenchida → pula
                    if fm.get("key", "").strip():
                        skipped.append(name)
                        continue

                    detected = _detect_key_py(body or content)
                    if not detected:
                        skipped.append(name)
                        continue

                    fm["key"] = detected
                    new_content = _build_md(fm, body)
                    update_md_content(svc, file_id, new_content)
                    fixed.append({"name": name, "key": detected})

                except Exception as e:
                    errors.append({"name": name, "error": str(e)})

    invalidate_library_cache()
    return jsonify({
        "fixed":   len(fixed),
        "skipped": len(skipped),
        "errors":  len(errors),
        "details": fixed,
    })


# ---------------------------------------------------------------------------
# YouTube Trending
# ---------------------------------------------------------------------------

_TRENDING_FILE = Path(__file__).parent / "_trending.json"
_TRENDING_MAX_AGE = 86400  # 24 horas

_CATHOLIC_ARTISTS_BR = [
    "Frei Gilson",
    "GBA Music",
    "Flavio Vitor Jr",
    "Ramon e Rafael",
    "Juninho Casimiro",
    "Rosa de Saron",
    "Colo de Deus",
    "Fraternidade São João Paulo II",
    "Missionário Shalom",
    "Comunidade Católica Shalom",
    "Anjos de Resgate",
    "Adriana Arydes",
    "Padre Marcelo Rossi",
    "Thiago Brado",
    "Livres Oficial",
    "Eliana Ribeiro",
    "Walmir Alencar",
    "Vida Reluz",
    "Voz da Verdade",
    "Frei Zezinho",
]

# Canais/artistas evangélicos que podem aparecer nos resultados — ignorar
_BLOCKED_CHANNELS = {
    "thiago brito", "gabriela rocha", "aline barros", "fernandinho",
    "diante do trono", "hillsong", "bethel", "elevation worship",
    "ana paula valadão", "eyshila", "isadora pompeo",
    "davi sacer", "midian lima", "bruna karla",
    "anderson freire", "kemuel", "thalles roberto",
}

_TITLE_BLOCK = {
    # Rezas e devoções — não são músicas
    "rosário", "terço", "novena", "via crucis", "hora santa",
    "adoração ao santíssimo", "quaresma", "| dia ",
    # Conteúdo não-musical
    "vlog", "reflexão", "meditação", "reel", "ep.", "podcast",
    "entrevista", "depoimento", "pregação", "homilia",
    # Compilações
    "1 hora", "2 horas", "3 horas", "4 horas", "5 horas",
    "coletânea", "compilado", "as melhores", "mix",
    "canto para missa", "cantos para missa", "canto de comunhão",
    "canto de entrada", "canto de ofertório", "hinário",
    "liturgia das horas", "hora de louvor",
}

_TITLE_ALLOW = {
    "cover", "ao vivo", "live", "letra", "clipe", "oficial",
    "louvor", "adoração", "canto", "canção", "hino", "música",
}

_TAGS_ALLOW = {"música", "música católica", "louvor", "worship", "canção", "hino"}

_MAX_DURATION_S = 5 * 60    # 5 minutos — apenas clipes/músicas curtas
_MAX_PER_CHANNEL = 2        # variedade: no máximo 2 vídeos por canal


def _parse_duration_seconds(iso: str) -> int:
    m = re.match(r"PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?", iso or "")
    if not m:
        return 0
    h, mn, s = (int(x or 0) for x in m.groups())
    return h * 3600 + mn * 60 + s


def _fmt_duration(iso: str) -> str:
    s = _parse_duration_seconds(iso)
    m, sec = divmod(s, 60)
    return f"{m}:{sec:02d}"


def _fmt_views(n: int) -> str:
    if n >= 1_000_000:
        v = n / 1_000_000
        return f"{v:.1f}M".replace(".0M", "M")
    if n >= 1_000:
        return f"{n // 1_000}k"
    return str(n)


def _is_music_video(item: dict) -> bool:
    snippet = item.get("snippet", {})
    content = item.get("contentDetails", {})
    title = snippet.get("title", "").lower()
    category = snippet.get("categoryId", "")
    tags = {t.lower() for t in (snippet.get("tags") or [])}
    duration_s = _parse_duration_seconds(content.get("duration", ""))

    if duration_s <= 60 or duration_s > _MAX_DURATION_S:
        return False
    if any(w in title for w in _TITLE_BLOCK):
        return False
    if category != "10" and not any(w in title for w in _TITLE_ALLOW):
        return False
    return category == "10" or bool(tags & _TAGS_ALLOW) or any(w in title for w in _TITLE_ALLOW)


def _fetch_youtube_trending():
    import requests as rq
    from datetime import timezone, datetime as dt, timedelta
    api_key = os.environ.get("YOUTUBE_API_KEY", "")
    if not api_key:
        log.warning("[trending] YOUTUBE_API_KEY não definida")
        return []

    published_after = (dt.now(timezone.utc) - timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%SZ")

    # Passo 1: coleta candidatos via search.list por artista (armazena channelId)
    candidates = {}
    for artist in _CATHOLIC_ARTISTS_BR:
        if len(candidates) >= 60:
            break
        try:
            r = rq.get(
                "https://www.googleapis.com/youtube/v3/search",
                params={
                    "part": "snippet",
                    "q": artist,
                    "type": "video",
                    "maxResults": 5,
                    "order": "viewCount",
                    "publishedAfter": published_after,
                    "regionCode": "BR",
                    "relevanceLanguage": "pt",
                    "key": api_key,
                },
                timeout=10,
            )
            r.raise_for_status()
            for item in r.json().get("items", []):
                vid_id = item["id"]["videoId"]
                if vid_id in candidates:
                    continue
                snip = item["snippet"]
                thumbs = snip.get("thumbnails", {})
                thumb = (thumbs.get("medium") or thumbs.get("high") or thumbs.get("default") or {}).get("url", "")
                candidates[vid_id] = {
                    "videoId": vid_id,
                    "title": snip.get("title", ""),
                    "channel": snip.get("channelTitle", ""),
                    "channelId": snip.get("channelId", ""),
                    "thumbnail": thumb,
                }
        except Exception as e:
            log.error("[trending] erro buscando '%s': %s", artist, e)

    if not candidates:
        return []

    # Passo 2a: verifica país dos canais — exclui canais com país explicitamente não-BR
    channel_ids = list({v["channelId"] for v in candidates.values() if v.get("channelId")})
    non_br_channels = set()
    for i in range(0, len(channel_ids), 50):
        batch = channel_ids[i:i + 50]
        try:
            r = rq.get(
                "https://www.googleapis.com/youtube/v3/channels",
                params={"part": "snippet", "id": ",".join(batch), "key": api_key},
                timeout=10,
            )
            r.raise_for_status()
            for ch in r.json().get("items", []):
                country = ch.get("snippet", {}).get("country", "").upper()
                if country and country != "BR":
                    non_br_channels.add(ch["id"])
        except Exception as e:
            log.error("[trending] erro em channels.list: %s", e)

    # Passo 2b: busca detalhes em batch (snippet + contentDetails + statistics)
    detailed = {}
    vid_ids = list(candidates.keys())
    for i in range(0, len(vid_ids), 50):
        batch = vid_ids[i:i + 50]
        try:
            r = rq.get(
                "https://www.googleapis.com/youtube/v3/videos",
                params={
                    "part": "snippet,contentDetails,statistics",
                    "id": ",".join(batch),
                    "key": api_key,
                },
                timeout=10,
            )
            r.raise_for_status()
            for item in r.json().get("items", []):
                detailed[item["id"]] = item
        except Exception as e:
            log.error("[trending] erro em videos.list: %s", e)

    # Passo 3 e 4: filtra, limita por canal e ordena por views
    channel_counts: dict = {}
    results = []
    for vid_id, candidate in candidates.items():
        if candidate.get("channelId") in non_br_channels:
            continue
        if candidate.get("channel", "").lower().strip() in _BLOCKED_CHANNELS:
            continue
        detail = detailed.get(vid_id)
        if not detail or not _is_music_video(detail):
            continue
        ch = candidate["channel"]
        if channel_counts.get(ch, 0) >= _MAX_PER_CHANNEL:
            continue
        channel_counts[ch] = channel_counts.get(ch, 0) + 1
        view_count = int(detail.get("statistics", {}).get("viewCount", 0))
        iso_dur = detail.get("contentDetails", {}).get("duration", "")
        candidate["viewCount"] = view_count
        candidate["views"] = _fmt_views(view_count)
        candidate["duration"] = _fmt_duration(iso_dur)
        results.append(candidate)

    results.sort(key=lambda v: v.get("viewCount", 0), reverse=True)
    for v in results:
        v.pop("viewCount", None)
        v.pop("channelId", None)
    return results[:10]


def _refresh_youtube_trending():
    try:
        if _TRENDING_FILE.exists():
            age = time.time() - _TRENDING_FILE.stat().st_mtime
            if age < _TRENDING_MAX_AGE:
                return
        log.info("[trending] atualizando cache YouTube...")
        data = _fetch_youtube_trending()
        if data:
            _TRENDING_FILE.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
            log.info("[trending] %d vídeos salvos", len(data))
    except Exception as e:
        log.error("[trending] falha no refresh: %s", e)


def _trending_worker():
    while True:
        _refresh_youtube_trending()
        time.sleep(3600)  # verifica a cada hora; busca só se cache > 24h

threading.Thread(target=_trending_worker, daemon=True).start()


@app.route("/api/trending")
@login_required
def api_trending():
    if _TRENDING_FILE.exists():
        return Response(_TRENDING_FILE.read_text(encoding="utf-8"), mimetype="application/json")
    return jsonify([])


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    mode = "Google Drive + OAuth" if _use_drive() else f"Local ({CIFRAS_ROOT})"
    auth = "OAuth ativo" if is_oauth_configured() else "sem autenticação"
    print(f"Modo  : {mode}")
    print(f"Auth  : {auth}")
    app.run(host="0.0.0.0", port=5000, debug=True, threaded=True)
